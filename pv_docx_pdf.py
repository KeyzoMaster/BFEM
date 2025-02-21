from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3 as sq

conn = sq.connect("BD2.db")
cur = conn.cursor()

document = Document('C:/Users/TBE/Documents/impression_pv.docx')

# TITRE de fichier
titre = document.add_heading("OFFICE du BFEEM", 1)
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

# les informations du jury et du centre d'examen
document.add_heading(' ', 2)
list_paragraph = document.add_paragraph()
list_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
list_paragraph.style = ''
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('IA:').bold = True
list_paragraph.add_run(input('Donner le nom de l\'IA'))
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('IEF:').bold = True
list_paragraph.add_run(input('Donner le nom de l\'IEF'))
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('Centre d\'Examen:').bold = True
list_paragraph.add_run(input('Donner le nom du centre d\'examen'))
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('JURY:').bold = True
list_paragraph.add_run(input('Donner le nom du jury'))
list_paragraph.add_run('\n')

# le pv de deliberation
document.add_heading('PV de DELIBERATION', 2)
document.add_heading('ADMIS D\'OFFICE AU BFEEM', 3)
cur.execute(
    """SELECT Numero_table,((Compo_Fran * Coef1) + (Dictee * Coef2) + (Etude_de_texte * Coef3) + (Instruction_Civique 
    * Coef4) +(Histoire_Geographie * Coef5) +(Mathematique * Coef6) +(PC_LV2 * Coef7) +(SVT * Coef8) +(Anglais1 * 
    Coef9) +(Anglais_Orale * Coef10) + CASE WHEN EPS > 10 THEN (EPS - 10) ELSE (10 - EPS) END) AS total_point FROM 
    Note inner join Candidat)""")
total = cur.fetchall()
table_Office = document.add_table(rows=1, cols=13)
table_Office.style = 'Table Grid'
document.add_heading('ADMIS AU 2ND TOUR', 4)
table_2tour = document.add_table(rows=1, cols=13)
table_2tour.style = 'Table Grid'
for row in total:
    for total_point in row:
        if total_point > 180:
            table_Office.add_row(row[total_point])
        elif 180 > total_point >= 153:
            table_2tour.add_row(row[total_point])
        else:
            print("")

document.save('impression_pv.docx')

# convertir le document .docx en .pdf

chemin_word = "impression_pv.docx"
chemin_pdf = "impression_pv.pdf"

convert(chemin_word, chemin_pdf)

conn.commit()
conn.close()
