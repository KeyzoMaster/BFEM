from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3 as sq
conn = sq.connect("BD2.db")
cur = conn.cursor()

document = Document('C:/Users/TBE/Documents/impression_candidat.docx')

# TITRE de fichier
titre = document.add_heading("OFFICE du BFEEM", 1)
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

# les informations du jury et du centre d'examen
document.add_heading(' ', 2)
list_paragraph = document.add_paragraph()
list_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
list_paragraph.style = ''
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('IA:').bold = True
list_paragraph.add_run(input('Donner le nom de l\'IA'))
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('IEF:').bold = True
list_paragraph.add_run(input('Donner le nom de l\'IEF'))
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('Centre d\'Examen:').bold = True
list_paragraph.add_run(input('Donner le nom du centre d\'examen'))
list_paragraph.add_run('\n')
list_paragraph.add_run('\n')
list_paragraph.add_run('JURY:').bold = True
list_paragraph.add_run(input('Donner le nom du jury'))
list_paragraph.add_run('\n')

# liste des candidats et les anonymats
titre2 = document.add_heading('LISTE DES CANDIDATS', 3)
titre2.alignment = WD_ALIGN_PARAGRAPH.CENTER
titre2.style = ''

table = document.add_table()
table.style = 'Table Grid'

cur.execute('SELECT * FROM Candidat')
table1 = cur.fetchall()
for row in table1:
    table.add_row(row)

# signature

signe = document.add_paragraph('Le responsable', 3)
signe.alignment = WD_ALIGN_PARAGRAPH.RIGHT

document.save('C:/Users/TBE/Documents/impression_candidat.docx')

# convertir le document .docx en document .pdf
chemin_word = "C:/Users/TBE/Documents/impression_candidat.docx"
chemin_pdf = "C:/Users/TBE/Documents/impression_candidat.pdf"

convert(chemin_word, chemin_pdf)

conn.commit()
conn.close()
