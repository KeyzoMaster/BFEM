from database import BrevetDB
from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ImpressBFEM:
    def __init__(self):
        jury_infos = BrevetDB().cursor.execute(
            """SELECT ia, ief, localite, centre_examen, president_jury, telephone FROM Jury""").fetchone()
        (self.ia, self.ief, self.localite, self.centre, self.president, self.tel) = jury_infos if jury_infos else (
            "", "", "", "", "", "")

    def document_header(self, document):
        document.add_heading(' ', 2)
        list_paragraph = document.add_paragraph()
        list_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        list_paragraph.style = None
        list_paragraph.add_run('\n')
        list_paragraph.add_run('\n')
        list_paragraph.add_run('IA:').bold = True
        list_paragraph.add_run(self.ia)
        list_paragraph.add_run('\n')
        list_paragraph.add_run('\n')
        list_paragraph.add_run('IEF:').bold = True
        list_paragraph.add_run(self.ief)
        list_paragraph.add_run('\n')
        list_paragraph.add_run('\n')
        list_paragraph.add_run('Centre d\'Examen:').bold = True
        list_paragraph.add_run(self.centre)
        list_paragraph.add_run('\n')
        list_paragraph.add_run('\n')
        list_paragraph.add_run('JURY:').bold = True
        list_paragraph.add_run(self.president)
        list_paragraph.add_run('\n')

    def liste_candidats(self):
        document = Document()

        # TITRE de fichier
        titre = document.add_heading("OFFICE du BFEM", 1)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # les informations du jury et du centre d'examen
        self.document_header(document)

        # liste des candidats et les anonymats
        titre2 = document.add_heading('LISTE DES CANDIDATS', 3)
        titre2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titre2.style = None

        table = document.add_table(1, 11)
        table.style = 'Table Grid'
        ligne = table.rows[0].cells
        table1 = BrevetDB().personal_information_candidat(53)
        i = 0
        for k in table1.keys():
            ligne[i].text = k
            i += 1
            if i > 10:
                break

        nums = BrevetDB().liste_num()

        for num in nums:
            cells = table.add_row().cells
            infos = BrevetDB().personal_information_candidat(num)
            i = 0
            for v in infos.values():
                if i > 10:
                    break
                cells[i].text = v
                i += 1

        # signature
        signe = document.add_paragraph('Le responsable', None)
        signe.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.save('impression_candidat.docx')

    def liste_anonymats(self):
        pass

    def releves_note(self):
        resultats = BrevetDB().resultats()
        for r in resultats:
            self.releve_note(int(r["Numéro de table"]))

        resultats = BrevetDB().resultats(tour="SECOND TOUR")
        for r in resultats:
            self.releve_note(int(r["Numéro de table"]))

    def releve_note(self, num_table):
        document = Document()

        # TITRE de fichier
        titre = document.add_heading("OFFICE DU BFEM", 1)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # les informations du jury et du centre d'examen
        self.document_header(document)

        infos = BrevetDB().personal_information_candidat(int(num_table))

        document.add_heading('', 2)
        list_paragraph = document.add_paragraph()
        list_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        list_paragraph.add_run('\n')
        list_paragraph.add_run('Numero de table:').bold = True
        list_paragraph.add_run(str(num_table))
        list_paragraph.add_run('\n')
        list_paragraph.add_run(infos["Prénom(s)"])
        list_paragraph.add_run('\t')
        list_paragraph.add_run(infos["Nom"])

        document.add_heading('RELEVE DE NOTE DU CANDIDAT', 3)

        notes = BrevetDB().notes_information_candidat(num_table)

        list_paragraph = document.add_paragraph()
        table = document.add_table(2, 12)
        table.style = 'Table Grid'
        ligne = table.rows[0].cells
        ligne[0].text = "Numéro de table"
        i = 1
        for k in notes.keys():
            ligne[i].text = k
            i += 1
            if i > 11:
                break
        ligne = table.rows[1].cells
        ligne[0].text = str(num_table)
        i = 1
        for v in notes.values():
            ligne[i].text = v
            i += 1
            if i > 11:
                break

        document.save('impression_pv' + str(num_table) + '.docx')

    def pv_bfem(self, tour="PREMIER TOUR"):
        # le pv de deliberation
        # cur.execute("""SELECT * FROM  Candidat""")
        # candidats = cur.fetchall()
        document = Document()

        self.document_header(document)

        document.add_heading('PV de DELIBERATION', 2)
        document.add_heading('ADMIS D\'OFFICE AU BFEM', 3)

        if tour == "PREMIER TOUR":
            resultats = BrevetDB().resultats(tour=tour)
            table_office = document.add_table(1, 12)
            table_office.style = 'Table Grid'
            ligne = table_office.rows[0].cells
            ligne[11].text = "Rang"
            i = 0

            for k in resultats[0].keys():
                ligne[i].text = k
                i += 1
                if i > 10:
                    break

            i = 1
            for r in resultats:
                if r["Statut"] != "Admis(e)":
                    continue
                j = 0
                cells = table_office.add_row().cells
                cells[11].text = str(i)
                i += 1
                for v in r.values():
                    cells[j].text = v
                    j += 1
                    if j > 10:
                        break

            document.add_heading('ADMIS AU 2ND TOUR', 4)
            table_2tour = document.add_table(1, 11)
            table_2tour.style = 'Table Grid'
            ligne = table_2tour.rows[0].cells
            i = 0
            for k in resultats[0].keys():
                ligne[i].text = k
                i += 1
                if i > 10:
                    break

            i = 1
            for r in resultats:
                if r["Statut"] == "Admis(e)":
                    continue
                j = 0
                cells = table_2tour.add_row().cells
                i += 1
                for v in r.values():
                    cells[j].text = v
                    j += 1
                    if j > 10:
                        break
        else:
            resultats = BrevetDB().resultats(tour=tour)
            table_office = document.add_table(1, 12)
            table_office.style = 'Table Grid'
            ligne = table_office.rows[0].cells
            ligne[11].text = "Rang"
            i = 0

            for k in resultats[0].keys():
                ligne[i].text = k
                i += 1
                if i > 11:
                    break

            i = 1
            for r in resultats:
                if r["Statut"] != "Admis(e)":
                    continue
                j = 0
                cells = table_office.add_row().cells
                cells[11].text = str(i)
                i += 1
                for v in r.values():
                    cells[j].text = v
                    j += 1
                    if j > 11:
                        break

        document.save('impression_pv.docx')

